# src/utils/export.py

import io
from datetime import datetime

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
)
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document


class ExportManager:

    # ----------------------------------------------------
    # Markdown
    # ----------------------------------------------------

    @staticmethod
    def markdown(text):

        return text

    # ----------------------------------------------------
    # TXT
    # ----------------------------------------------------

    @staticmethod
    def text(text):

        return text

    # ----------------------------------------------------
    # PDF
    # ----------------------------------------------------

    @staticmethod
    def pdf(title, text):

        buffer = io.BytesIO()

        document = SimpleDocTemplate(buffer)

        styles = getSampleStyleSheet()

        story = []

        story.append(
            Paragraph(title, styles["Title"])
        )

        story.append(
            Paragraph(
                f"Generated: {datetime.now()}",
                styles["Normal"],
            )
        )

        story.append(
            Paragraph("<br/><br/>", styles["Normal"])
        )

        for line in text.split("\n"):

            if line.strip():

                story.append(
                    Paragraph(
                        line.replace("<", "&lt;").replace(">", "&gt;"),
                        styles["BodyText"],
                    )
                )

        document.build(story)

        pdf = buffer.getvalue()

        buffer.close()

        return pdf

    # ----------------------------------------------------
    # DOCX
    # ----------------------------------------------------

    @staticmethod
    def docx(title, text):

        document = Document()

        document.add_heading(title, level=1)

        document.add_paragraph(
            f"Generated: {datetime.now()}"
        )

        document.add_paragraph("")

        for line in text.split("\n"):

            document.add_paragraph(line)

        output = io.BytesIO()

        document.save(output)

        output.seek(0)

        return output.getvalue()

    # ----------------------------------------------------
    # File Name
    # ----------------------------------------------------

    @staticmethod
    def filename(prefix, extension):

        timestamp = datetime.now().strftime(
            "%Y%m%d_%H%M%S"
        )

        return f"{prefix}_{timestamp}.{extension}"